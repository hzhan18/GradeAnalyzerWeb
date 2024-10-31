# report_generation.py

import logging
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ai_interface import call_with_messages
from docx.oxml.ns import qn  # 用于支持中文字体设置
from docx.enum.section import WD_SECTION


def set_paragraph_font(paragraph, size=10.5, name='SimSun'):
    """
    设置段落中所有Run对象的字体大小和字体名称。
    """
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.name = name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), name)  # 确保中文字体设置

def add_logo_to_header(doc, logo_path, width_cm):
    """
    在文档的页眉左上角添加一个图标。

    参数：
    - doc: Document对象
    - logo_path: 图标的文件路径
    - width_cm: 图标的宽度（厘米）
    """
    # 遍历所有sections，添加页眉
    for section in doc.sections:
        header = section.header
        # 添加一个段落
        paragraph = header.paragraphs[0]
        # 清除默认内容
        paragraph.clear()
        # 设置段落左对齐
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 添加图标
        run = paragraph.add_run()
        try:
            run.add_picture(logo_path, width=Cm(width_cm))
        except Exception as e:
            logging.error(f"无法添加页眉图标: {e}")
            run.add_text("图标加载失败")

def generate_score_table(doc, score_data, total_students, table_title, table_number):
    """
    生成成绩分析表格，包括表格标题、表格内容和分析结果。

    参数：
    - doc: Document对象
    - score_data: 单项成绩的数据字典，包括 'stats', 'distribution_text', 'plot_file_name' 等
    - total_students: 总人数
    - table_title: 表格标题，如 "网络学习部分成绩情况表"
    - table_number: 表格编号，如 "表2"
    """
    # 添加表格标题段落
    table_title_paragraph = doc.add_paragraph()
    table_title_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
    table_title_run = table_title_paragraph.add_run(f"{table_number}  {table_title}")
    table_title_run.font.size = Pt(9)
    table_title_run.font.name = 'SimSun'
    # 加粗表格编号
    for run in table_title_paragraph.runs:
        if table_number in run.text:
            run.bold = True

    # 创建表格
    table = doc.add_table(rows=5, cols=11)
    table.style = 'Table Grid'
    table.autofit = False
    widths = [Inches(0.6)] * 10 + [Inches(1.0)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # 填充表格内容
    # 第一行
    table.cell(0, 0).text = "人数："
    table.cell(0, 1).text = str(total_students)

    # 合并第三列和第四列，写"最高分"
    highest_score_cell = table.cell(0, 2).merge(table.cell(0, 3))
    highest_score_cell.text = "最高分:"
    highest_score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(highest_score_cell.paragraphs[0], size=9, name='SimSun')

    # 第五列填写最高分的数值
    if score_data:
        table.cell(0, 4).text = str(score_data['stats']['最高分'])
    else:
        table.cell(0, 4).text = "N/A"

    # 合并第六列和第七列，写"最低分："
    lowest_score_cell = table.cell(0, 5).merge(table.cell(0, 6))
    lowest_score_cell.text = "最低分："
    lowest_score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(lowest_score_cell.paragraphs[0], size=9, name='SimSun')

    # 第八列填写最低分的数值
    if score_data:
        table.cell(0, 7).text = str(score_data['stats']['最低分'])
    else:
        table.cell(0, 7).text = "N/A"

    # 合并第九列和第十列，写"平均分："
    average_score_cell = table.cell(0, 8).merge(table.cell(0, 9))
    average_score_cell.text = "平均分："
    average_score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(average_score_cell.paragraphs[0], size=9, name='SimSun')

    # 第十一列填写平均分的数值
    if score_data:
        table.cell(0, 10).text = str(score_data['stats']['平均分'])
    else:
        table.cell(0, 10).text = "N/A"

    # 第二行
    table.cell(1, 0).text = "分数段"
    score_segments = ["0-9", "10-19", "20-29", "30-39", "40-49", "50-59", "60-69", "70-79", "80-89", "90-100"]
    for idx, segment in enumerate(score_segments):
        if idx + 1 < len(table.columns):
            table.cell(1, idx + 1).text = segment
        else:
            break

    # 第三行：人数
    table.cell(2, 0).text = "人数"
    if score_data:
        for idx, (segment, values) in enumerate(score_data['distribution_text'].items()):
            if idx + 1 < len(table.columns):
                count = values.get('人数', 0)
                table.cell(2, idx + 1).text = str(count)
            else:
                break
    else:
        for idx in range(1, len(table.columns)):
            table.cell(2, idx).text = "N/A"

    # 第四行：比例
    table.cell(3, 0).text = "比例"
    if score_data:
        for idx, (segment, values) in enumerate(score_data['distribution_text'].items()):
            if idx + 1 < len(table.columns):
                percentage = values.get('占比', 0.0)
                table.cell(3, idx + 1).text = f"{percentage:.2f}%"
            else:
                break
    else:
        for idx in range(1, len(table.columns)):
            table.cell(3, idx).text = "N/A"

    # 第五行：成绩分布图
    table.cell(4, 0).text = "成绩分布图"
    # 合并第二列到第十一列
    row5 = table.rows[4]
    row5.cells[1].merge(row5.cells[10])
    if score_data and 'plot_file_name' in score_data:
        try:
            run = table.cell(4, 1).paragraphs[0].add_run()
            run.add_picture(score_data['plot_file_name'], width=Inches(6))
        except Exception as e:
            table.cell(4, 1).text = "图片插入失败"
    else:
        table.cell(4, 1).text = "N/A"

    # 设置表格字体和居中
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if idx < 3:
                    set_paragraph_font(paragraph, size=9, name='SimSun')
                else:
                    set_paragraph_font(paragraph, size=10.5, name='SimSun')

def generate_word_report(
    base_name,
    output_file_name,
    all_scores_data,
    output_path,
    semester_info,
    course_name,
    total_students,
    class_name1,
    class_name2
):  
    try:
        logging.info("Generating report document...")
        # Define output_file_name with full path in 'uploads' folder
        output_file_name = os.path.join("uploads", base_name, f"{base_name}_Summary.docx")
        
        # Ensure the directory exists before saving the document
        output_dir = os.path.dirname(output_file_name)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            logging.info(f"Directory created: {output_dir}")
        else:
            logging.info(f"Directory already exists: {output_dir}")

        # Check permissions for writing the document
        if os.access(output_dir, os.W_OK):
            logging.info("Directory is writable")
        else:
            logging.error("Directory is not writable")

        # Attempt to save a dummy file in the directory to confirm write access
        dummy_path = os.path.join(output_dir, "dummy_test_file.txt")
        try:
            with open(dummy_path, 'w') as dummy_file:
                dummy_file.write("This is a test.")
            logging.info(f"Dummy file successfully created at: {dummy_path}")
            os.remove(dummy_path)  # Clean up the dummy file
        except Exception as e:
            logging.error(f"Failed to create a dummy file in directory: {output_dir} - Error: {e}")
            return None  # Exit if we can't write to the directory

        # 设置中文字体支持
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        # 设置页边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.2)       # 上边距2.2厘米
            section.bottom_margin = Cm(1.37)   # 下边距1.37厘米
            section.left_margin = Cm(2)         # 左边距2厘米
            section.right_margin = Cm(2)        # 右边距2厘米

        # 获取当前脚本文件的绝对路径
        script_dir = os.path.dirname(os.path.abspath(__file__))
        # 构建logo.png的绝对路径
        logo_path = os.path.join(script_dir, "logo.png")
        # 添加页眉图标
        add_logo_to_header(doc, logo_path, width_cm=3.5)

        # 添加标题：居中加粗18号SimHei，写着"课程反思报告"
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("课程反思报告")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.name = 'SimHei'
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')  # 确保SimHei字体用于中文

        # 第二行：居中Times New Roman 10.5号字，内容是学期信息
        semester_paragraph = doc.add_paragraph()
        semester_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        semester_run = semester_paragraph.add_run('('+semester_info+')')
        semester_run.font.size = Pt(10.5)
        semester_run.font.name = 'Times New Roman'

        # 第三行：加粗10.5号SimSun字，写"一、教学基本信息"
        basic_info_title = doc.add_paragraph()
        basic_info_title_run = basic_info_title.add_run("一、教学基本信息")
        basic_info_title_run.bold = True
        basic_info_title_run.font.size = Pt(10.5)
        basic_info_title_run.font.name = 'SimSun'
        basic_info_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 第四行：缩进两个字符，10.5号SimSun字，写"课程名称：" + course_name
        course_paragraph = doc.add_paragraph()
        course_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符（约21pt）
        course_run = course_paragraph.add_run(f"课程名称：{course_name}")
        set_paragraph_font(course_paragraph, size=10.5, name='SimSun')

        # 第五行：缩进两个字符，10.5号SimSun字，写"授课对象及合班情况：" + total_students
        if class_name1 and class_name2:
            class_info = f"{class_name1}和{class_name2}，共计{total_students}人"
        elif class_name1:
            class_info = f"{class_name1}，共计{total_students}人"
        else:
            class_info = f"共计{total_students}人"  # 如果没有输入班级名称
        
        student_paragraph = doc.add_paragraph()
        student_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        student_run = student_paragraph.add_run(f"授课对象及合班情况：{class_info}")
        set_paragraph_font(student_paragraph, size=10.5, name='SimSun')

        prof_paragraph = doc.add_paragraph()
        prof_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        prof_run = prof_paragraph.add_run(f"任课教师：刘利钊")
        set_paragraph_font(prof_paragraph, size=10.5, name='SimSun')

        sem_time_paragraph = doc.add_paragraph()
        sem_time_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        prof_run = sem_time_paragraph.add_run(f"实际授课课时：24学时")
        set_paragraph_font(sem_time_paragraph, size=10.5, name='SimSun')

        # 第六行：加粗10.5号SimSun字，写"二、考核及学习成绩情况"
        exam_info_title = doc.add_paragraph()
        exam_info_title_run = exam_info_title.add_run("二、考核及学习成绩情况")
        exam_info_title_run.bold = True
        exam_info_title_run.font.size = Pt(10.5)
        exam_info_title_run.font.name = 'SimSun'
        exam_info_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 添加接下来的内容
        # 1. 10.5号SimSun字，每段缩进两个字符
        content_paragraph = doc.add_paragraph()
        content_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        content_run = content_paragraph.add_run(
            "本学期的《大学信息技术》课程采用SPOC教学模式，理论课选用国家精品线上课程（福建农林大学陈琼老师的《大学信息技术基础》）为蓝本，要求学生通过网课平台进行自学，实验课则安排在实验室进行线下授课。"
        )
        set_paragraph_font(content_paragraph, size=10.5, name='SimSun')

        content_paragraph = doc.add_paragraph()
        content_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        content_run = content_paragraph.add_run(
            "课程的总评成绩计算方式如下："
        )
        set_paragraph_font(content_paragraph, size=10.5, name='SimSun')

        content_paragraph = doc.add_paragraph()
        content_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        content_run = content_paragraph.add_run(
            "总评成绩=期末考试成绩×60%+实验成绩×20%+网络学习成绩×20%"
        )
        set_paragraph_font(content_paragraph, size=10.5, name='SimSun')

        content_paragraph = doc.add_paragraph()
        content_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        content_run = content_paragraph.add_run(
            "特别要求：当期末考试成绩≥50分时，可直接按上述公式计算总评成绩；当期末考试成绩＜50分时，则在上述公式结果和期末考试成绩之间选择最低值作为总评成绩。"
        )
        set_paragraph_font(content_paragraph, size=10.5, name='SimSun')

        # 2. 10.5号SimSun字加粗，写"1、网络学习部分成绩的计算"
        section1_title = doc.add_paragraph()
        section1_title_run = section1_title.add_run("1、网络学习部分成绩的计算")
        section1_title_run.bold = True
        section1_title_run.font.size = Pt(10.5)
        section1_title_run.font.name = 'SimSun'
        section1_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 3. 10.5号SimSun字，缩进两个字符
        scoring_details_paragraph = doc.add_paragraph()
        scoring_details_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        scoring_details_run = scoring_details_paragraph.add_run("本学期网络学习成绩的评分细则如表1所示：")
        set_paragraph_font(scoring_details_paragraph, size=10.5, name='SimSun')

        # 4. 9号SimSun字，缩进两个字符，其中“表1”两字加粗
        table1_title_paragraph = doc.add_paragraph()
        table1_title_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符
        table1_title_run = table1_title_paragraph.add_run("表1  网络学习成绩评分细则表")
        table1_title_run.font.size = Pt(9)
        table1_title_run.font.name = 'SimSun'
        # 加粗“表1”
        for run in table1_title_paragraph.runs:
            if "表1" in run.text:
                run.bold = True

        # 5. 插入4行3列的表格，字居中，9号SimSun字
        table1 = doc.add_table(rows=4, cols=3)
        table1.style = 'Table Grid'
        # 设置表格宽度
        table1.autofit = False
        widths = [Inches(1.5), Inches(1.5), Inches(3)]
        for row in table1.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        # 填充表格内容
        table1.cell(0, 0).text = "评分内容"
        table1.cell(0, 1).text = "权重"
        table1.cell(0, 2).text = "说明"

        table1.cell(1, 0).text = "单元测验"
        table1.cell(1, 1).text = "60%"
        table1.cell(1, 2).text = "每次测验有3次机会（随机抽题），取最高分"

        table1.cell(2, 0).text = "综合测验"
        table1.cell(2, 1).text = "30%"
        table1.cell(2, 2).text = "只有1次机会（随机抽题）"

        table1.cell(3, 0).text = "讨论"
        table1.cell(3, 1).text = "10%"
        table1.cell(3, 2).text = "10次得满分，限课堂交流区，其它地方的讨论不计入成绩"

        # 设置表格字体和居中
        for row in table1.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'SimSun'

        # 添加表格和分析内容
        table_number = 2  # 表格编号从2开始，因为表1已经在之前添加

        for idx, score_data in enumerate(all_scores_data):
            score_type = score_data['score_type']
            if class_name1 and class_name2:
                table2_class_info = f"{class_name1}和{class_name2}"
            elif class_name1:
                table2_class_info = f"{class_name1}"
            else:
                table2_class_info = f"XX"  # 如果没有输入班级名称

            if score_type == "平时":
                # 添加网络学习部分的表格和内容
                table_title = "网络学习部分成绩情况表"
                # 在添加表格前，添加提示语句
                table2_reference_paragraph = doc.add_paragraph()
                table2_reference_paragraph.paragraph_format.first_line_indent = Pt(21)  # 缩进两个字符

                table2_reference_run = table2_reference_paragraph.add_run(f"{table2_class_info}班的网络学习部分的成绩情况见表2。")
                set_paragraph_font(table2_reference_paragraph, size=10.5, name='SimSun')
            elif score_type == "实验":
                # 添加实验成绩部分的提示和标题
                # 加粗的10.5号SimSun字，写上“2、实验成绩的计算”
                section_title = doc.add_paragraph()
                section_title_run = section_title.add_run("2、实验成绩的计算")
                section_title_run.bold = True
                section_title_run.font.size = Pt(10.5)
                section_title_run.font.name = 'SimSun'
                section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 10.5号SimSun字，缩进两个字符，写上提示语
                experiment_paragraph = doc.add_paragraph()
                experiment_paragraph.paragraph_format.left_indent = Pt(21)
                experiment_paragraph.add_run("本学期共进行了XX次实验，全部在万维考试系统上完成，由系统自动评分。\n"
                                            f"{table2_class_info}班的实验成绩情况见表3。")
                set_paragraph_font(experiment_paragraph, size=10.5, name='SimSun')

                table_title = "实验部分成绩情况表"
            elif score_type == "期末":
                # 添加期末考试成绩部分的提示和标题
                # 加粗的10.5号SimSun字，写上“3、期末考试成绩的计算”
                section_title = doc.add_paragraph()
                section_title_run = section_title.add_run("3、期末考试成绩的计算")
                section_title_run.bold = True
                section_title_run.font.size = Pt(10.5)
                section_title_run.font.name = 'SimSun'
                section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 10.5号SimSun字，缩进两个字符，写上提示语
                final_exam_paragraph = doc.add_paragraph()
                final_exam_paragraph.paragraph_format.first_line_indent = Pt(21)
                final_exam_paragraph.add_run(
                    "由于理论部分的教学内容已在网络课程中安排了单元测验和综合测验，因此期末考试主要考核实验部分的教学内容。本课程期末考试采用万维考试系统进行上机考试，考生登录考试系统后，系统自动为各考生分配试卷，每份试卷的题型和分值组成如下："
                )
                set_paragraph_font(final_exam_paragraph, size=10.5, name='SimSun')

                final_exam_paragraph = doc.add_paragraph()
                final_exam_paragraph.paragraph_format.left_indent = Pt(42)
                final_exam_paragraph.add_run(
                    "1) Windows/网络设置操作题15分。\n"
                    "2) MS Office Word操作题30分。\n"
                    "3) MS Office Excel操作题35分。\n"
                    "4) MS Office PPT操作题20分。"
                )
                set_paragraph_font(final_exam_paragraph, size=10.5, name='SimSun')

                final_exam_paragraph = doc.add_paragraph()
                final_exam_paragraph.paragraph_format.left_indent = Pt(21)
                final_exam_paragraph.add_run(
                    "从试题内容来看，试卷基本能覆盖教学大纲中的实践操作内容，能较好地考察学生的掌握情况，也能较真实地反映学生的计算机应用水平。\n"
                    f"{table2_class_info}班的期末考试成绩情况见表4。"
                )
                set_paragraph_font(final_exam_paragraph, size=10.5, name='SimSun')

                table_title = "期末考试成绩情况表"
            elif score_type == "总评":
                # 添加课程总评成绩部分的提示和标题
                # 加粗的10.5号SimSun字，写上“4、课程总评成绩情况”
                section_title = doc.add_paragraph()
                section_title_run = section_title.add_run("4、课程总评成绩情况")
                section_title_run.bold = True
                section_title_run.font.size = Pt(10.5)
                section_title_run.font.name = 'SimSun'
                section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 10.5号SimSun字，缩进两个字符，写上提示语
                total_score_paragraph = doc.add_paragraph()
                total_score_paragraph.paragraph_format.first_line_indent = Pt(21)
                total_score_paragraph.add_run(f"{table2_class_info}班的课程总评成绩情况见表5。")
                set_paragraph_font(total_score_paragraph, size=10.5, name='SimSun')

                table_title = "课程总评成绩情况表"
            else:
                table_title = f"{score_type}成绩情况表"

            # 调用生成表格的函数
            generate_score_table(
                doc=doc,
                score_data=score_data,
                total_students=total_students,
                table_title=table_title,
                table_number=f"表{table_number}"
            )

            # 添加成绩分析内容
            # 调用 AI 接口生成分析报告
            stats = score_data['stats']
            distribution_text = score_data['distribution_text']
            distribution_description = ""
            for range_key, values in distribution_text.items():
                distribution_description += f"{range_key}分数段: 人数 {values['人数']}, 占比 {values['占比']:.2f}%；"

            analysis_content = (
                f"你作为这门课的授课老师，正在写校方布置的课程总结报告，对该门课学生{score_type}的成绩做出简要的书面总结和分析(除非学生成绩数据比较特殊，否则请不要过多的展示各分数段的总结和过分的罗列数值，而是稍微宏观一些做出总结)。"
                f"以下为学生的成绩数据: 总人数为{stats['总人数']}，最高分为{stats['最高分']}，最低分为{stats['最低分']}，平均分为{stats['平均分']}。"
                f"各分数段的分布情况如下：{distribution_description}"
                "备注：生成的内容中要大幅减少转接词的使用（例如首先、其次、最后、综上所述、总的来说、此外、值得XX的是、XXXX的是）。"
                "同时要保证生成的内容通俗易懂，不晦涩，不要用太书面化的词语。将内容控制在500字以内。"
            )
            ai_result = call_with_messages(analysis_content)

            # 添加分析结果
            analysis_heading = doc.add_paragraph()
            analysis_heading_run = analysis_heading.add_run("分析结果:")
            analysis_heading_run.bold = True
            analysis_heading_run.font.size = Pt(10.5)
            analysis_heading_run.font.name = 'SimSun'
            analysis_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            analysis_paragraph = doc.add_paragraph(ai_result)
            analysis_paragraph.paragraph_format.left_indent = Pt(21)
            set_paragraph_font(analysis_paragraph, size=10.5, name='SimSun')

            table_number += 1  # 更新表格编号

            progress["value"] += 5
            status_label.config(text=f"正在生成 {score_type} 成绩报告...")
            root.update_idletasks()

        # 添加学习成效分析
        # 加粗的10.5号SimSun字，写上“三、学习成效分析”
        learning_heading = doc.add_paragraph()
        learning_heading_run = learning_heading.add_run("三、学习成效分析")
        learning_heading_run.bold = True
        learning_heading_run.font.size = Pt(10.5)
        learning_heading_run.font.name = 'SimSun'
        learning_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 10.5号SimSun字，缩进两个字符，写上提示语
        learning_intro_paragraph = doc.add_paragraph()
        learning_intro_paragraph.paragraph_format.left_indent = Pt(21)
        learning_intro_paragraph.add_run(
            "从学生的学习过程情况及期末总评成绩来看，大部分学生均能完成本课程的学习任务，其计算机基础知识和计算机应用水平基本达到本课程的教学要求，本课程的教学目的基本实现。"
        )
        set_paragraph_font(learning_intro_paragraph, size=10.5, name='SimSun')

        # 调用 AI 接口生成学习成效分析
        learning_effectiveness_content = (
            "（然后从下面的方面进行分析："
            "1、学习本课程之前学生的课程基础的掌握情况；"
            "2、学习本课程时，学生的学习态度情况（如提问情况、主动性）；"
            "3、课时安排对教学效果的影响；"
            "4、教学方式对教学效果的影响；"
            "5、教学内容对教学效果的影响。）"
            "请注意，分析内容不要过于细节，可以适当使用较为概括和笼统的语言。"
            "将内容控制在500字以内。"
        )
        ai_learning_effectiveness_result = call_with_messages(learning_effectiveness_content)

        # 添加AI生成的学习成效分析
        cleaned_ai_learning_effectiveness_result = ai_learning_effectiveness_result.replace("#", "").replace("*", "")
        learning_paragraph = doc.add_paragraph(cleaned_ai_learning_effectiveness_result)
        learning_paragraph.paragraph_format.left_indent = Pt(21)
        set_paragraph_font(learning_paragraph, size=10.5, name='SimSun')

        # 添加改进措施及建议
        # 加粗的10.5号SimSun字，写上“四、改进措施及建议”
        suggestion_heading = doc.add_paragraph()
        suggestion_heading_run = suggestion_heading.add_run("四、改进措施及建议")
        suggestion_heading_run.bold = True
        suggestion_heading_run.font.size = Pt(10.5)
        suggestion_heading_run.font.name = 'SimSun'
        suggestion_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 调用 AI 接口生成改进措施及建议
        suggestion_content = (
            "请写出几条针对这个课程方方面面都可以的改进措施及建议。"
            "请注意，分析内容不要过于细节，可以适当使用较为概括和笼统的语言。"
            "将内容控制在500字以内。"
        )
        ai_suggestion_result = call_with_messages(suggestion_content)

        # 添加AI生成的改进措施及建议
        cleaned_ai_suggestion_result = ai_suggestion_result.replace("#", "").replace("*", "")
        suggestion_paragraph = doc.add_paragraph(cleaned_ai_suggestion_result)
        suggestion_paragraph.paragraph_format.left_indent = Pt(21)
        set_paragraph_font(suggestion_paragraph, size=10.5, name='SimSun')
        print(f"Saving document to: {output_file_name}")

        # Ensure the directory exists before saving the document
        output_dir = os.path.dirname(output_file_name)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            logging.info(f"Directory created: {output_dir}")
        else:
            logging.info(f"Directory already exists: {output_dir}")

        # Check permissions for writing the document
        if os.access(output_dir, os.W_OK):
            logging.info("Directory is writable")
        else:
            logging.error("Directory is not writable")

        # Attempt to save a dummy file in the directory to confirm write access
        dummy_path = os.path.join(output_dir, "dummy_test_file.txt")
        try:
            with open(dummy_path, 'w') as dummy_file:
                dummy_file.write("This is a test.")
            logging.info(f"Dummy file successfully created at: {dummy_path}")
        except Exception as e:
            logging.error(f"Failed to create a dummy file in directory: {output_dir} - Error: {e}")

        # Save the actual document
        logging.info(f"Saving document to: {output_file_name}")
        doc.save(output_file_name)

        # Confirm if the file is saved successfully
        if os.path.isfile(output_file_name):
            logging.info("Document saved successfully.")
        else:
            logging.error("Document file not found after save attempt: %s", output_file_name)
            return None

        return output_file_name if os.path.isfile(output_file_name) else None
    except Exception as e:
        logging.error("An error occurred while generating the Word report: %s", e)
        return None